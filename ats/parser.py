"""Extract plain text from resume/JD files (PDF, DOCX, TXT) or raw bytes."""

from __future__ import annotations

import io
from pathlib import Path


class UnsupportedFileType(ValueError):
    """Raised when a file extension is not supported."""


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts)


def _read_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    # Pull text out of tables too — resumes often use them for layout.
    for table in doc.tables:
        for row in table.rows:
            lines.extend(cell.text for cell in row.cells)
    return "\n".join(lines)


def _read_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_txt,
    ".text": _read_txt,
    ".md": _read_txt,
}

SUPPORTED_EXTENSIONS = tuple(_READERS)


def extract_text(data: bytes, filename: str) -> str:
    """Extract text from file bytes, dispatching on the filename extension."""
    ext = Path(filename).suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    return reader(data)


def extract_text_from_path(path: str | Path) -> str:
    """Read a file from disk and extract its text."""
    path = Path(path)
    return extract_text(path.read_bytes(), path.name)
